import docx                                      # Импортируем библиотеку "python docx"
import datetime                                  # Импортируем библиотеку "datetime"

dt = datetime.datetime.now()                     # Получаем текущую дату
tec_date = dt.strftime('%d.%m.%Y')               # Редактируем формат вывода текущий даты


# Блок авторизации...
maill = "osoben1@yandex.ru"                      # Данные для авторизации модератора login
password = "Вка5248"                             # Данные для авторизации password

# Блок управления контентом...
document = docx.Document('test.docx')            # Открываем тестовый документ "test.docx"

content_title = document.paragraphs[0].text      # Объявляем переменную для чтения 1 го параграфа
content_desc = document.paragraphs[1].text           # Объявляем переменную для чтения 2 го параграфа
content2 = document.paragraphs[2].text           # Объявляем переменную для чтения 3 го параграфа





